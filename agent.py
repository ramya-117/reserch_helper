import os
import re
import base64
import requests
from typing import TypedDict, List
from langgraph.graph import StateGraph, END
from langchain_openai import ChatOpenAI
from docx import Document
from resend import Resend as ResendClient

GROQ_API_KEY   = os.environ.get("GROQ_API_KEY")
EMAIL_SENDER   = os.environ.get("EMAIL_SENDER")
RESEND_API_KEY = os.environ.get("RESEND_API_KEY")

llm = ChatOpenAI(
    model="llama-3.3-70b-versatile",
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1",
    temperature=0.3,
)

class ResearchState(TypedDict):
    topic: str
    receiver_email: str
    papers: List[dict]
    summaries: List[str]
    outline: str
    draft: str
    references: List[str]
    email_body: str
    log: callable

def clean_abstract(text: str) -> str:
    if not text:
        return "Abstract not available."
    return re.sub("<.*?>", "", text)

def search_published_papers(state: ResearchState):
    state["log"](f"🔍 Searching papers for: {state['topic']}")
    url = "https://api.crossref.org/works"
    params = {"query": state["topic"], "rows": 5, "sort": "relevance"}
    try:
        response = requests.get(url, params=params, timeout=15)
        items = response.json().get("message", {}).get("items", [])
    except Exception as e:
        state["log"](f"❌ Search failed: {e}")
        return {"papers": [], "references": []}

    papers, references = [], []
    for paper in items:
        title    = paper.get("title", ["Untitled"])[0]
        year     = paper.get("issued", {}).get("date-parts", [[None]])[0][0]
        journal  = paper.get("container-title", ["Unknown Journal"])[0] if paper.get("container-title") else "Unknown Journal"
        doi      = paper.get("DOI", "")
        abstract = clean_abstract(paper.get("abstract", ""))
        authors  = ", ".join([f"{a.get('family','')} {a.get('given','')}" for a in paper.get("author", [])]) or "Unknown Author"
        papers.append({"title": title, "abstract": abstract, "year": year, "journal": journal, "authors": authors, "doi": doi})
        references.append(f"{authors} ({year}). {title}. {journal}. https://doi.org/{doi}")

    state["log"](f"✅ Found {len(papers)} papers")
    return {"papers": papers, "references": references}

def summarize_papers(state: ResearchState):
    summaries = []
    for i, paper in enumerate(state["papers"]):
        state["log"](f"📖 Summarizing paper {i+1}/{len(state['papers'])}: {paper['title'][:50]}...")
        prompt = f"""You are an academic research assistant. Summarize the following research paper.
Structure:
1. KEY CONTRIBUTION
2. METHODOLOGY
3. KEY FINDINGS
4. LIMITATIONS

Title: {paper['title']}
Authors: {paper['authors']}
Journal: {paper['journal']} ({paper['year']})
Abstract: {paper['abstract']}"""
        response = llm.invoke(prompt)
        summaries.append(f"Paper: {paper['title']}\n{response.content}")
    return {"summaries": summaries}

def generate_outline(state: ResearchState):
    state["log"]("📝 Generating research outline...")
    summaries_text = "\n\n---\n\n".join(state["summaries"])
    prompt = f"""Create a structured research paper outline.
Topic: {state['topic']}
Literature: {summaries_text}
Include: Introduction, Literature Review, Methodology, Results, Conclusion, Future Work"""
    response = llm.invoke(prompt)
    return {"outline": response.content}

def write_draft(state: ResearchState):
    state["log"]("✍️ Writing full draft (this takes ~1 min)...")
    prompt = f"""Write a complete formal research paper.
Topic: {state['topic']}
Outline: {state['outline']}
- Use APA citations, formal tone, full paragraphs, 1000+ words, include abstract."""
    response = llm.invoke(prompt)
    return {"draft": response.content}

def export_doc(state: ResearchState):
    state["log"]("💾 Creating DOCX file...")
    doc = Document()
    doc.add_heading(state["topic"], level=1)
    doc.add_heading("Research Paper", level=2)
    doc.add_paragraph(state["draft"])
    doc.add_heading("Literature Review Summaries", level=2)
    for summary in state["summaries"]:
        doc.add_paragraph(summary)
        doc.add_paragraph("")
    doc.add_heading("References", level=2)
    for ref in state["references"]:
        doc.add_paragraph(ref)
    doc.save("published_research_paper.docx")
    return {}

def prepare_email(state: ResearchState):
    papers, summaries = state["papers"], state["summaries"]
    rows = ""
    for i, (paper, summary) in enumerate(zip(papers, summaries)):
        contribution = next((l.replace("1.", "").replace("KEY CONTRIBUTION:", "").strip()
                             for l in summary.split("\n") if "KEY CONTRIBUTION" in l.upper()), "")
        bg = '#f9f9f9' if i % 2 == 0 else '#ffffff'
        rows += f"""<tr style="background:{bg}">
            <td style="padding:12px;font-weight:bold;color:#2c3e50;">[{i+1}] {paper['title']}</td></tr>
        <tr style="background:{bg}"><td style="padding:0 12px 12px 12px;color:#555;font-size:13px;">
            👤 <b>Authors:</b> {paper['authors']}<br>
            📰 <b>Journal:</b> {paper['journal']} ({paper['year']})<br>
            🔗 <b>DOI:</b> <a href="https://doi.org/{paper['doi']}">https://doi.org/{paper['doi']}</a><br>
            💡 <b>Key Contribution:</b> {contribution}</td></tr>"""

    html = f"""<html><body style="font-family:Arial,sans-serif;max-width:700px;margin:auto">
    <div style="background:#2c3e50;color:white;padding:20px;border-radius:8px 8px 0 0">
        <h1 style="margin:0">🔬 AI Research Agent Report</h1></div>
    <div style="background:#ecf0f1;padding:15px">
        <b>📌 Topic:</b> {state['topic']}<br><b>📚 Papers:</b> {len(papers)}</div>
    <h2 style="color:#2c3e50">📄 Paper Summaries</h2>
    <table width="100%" style="border:1px solid #ddd">{rows}</table>
    <h2 style="color:#2c3e50">📋 Outline</h2>
    <div style="background:#f4f4f4;padding:15px;white-space:pre-wrap;font-size:13px">{state['outline'][:1500]}...</div>
    <div style="background:#2ecc71;color:white;padding:15px;border-radius:6px;text-align:center">
        ✅ Full paper attached as <b>published_research_paper.docx</b></div>
    </body></html>"""
    return {"email_body": html}

def send_email(state: ResearchState):
    state["log"](f"📤 Sending email to {state['receiver_email']}...")
    try:
        if not RESEND_API_KEY:
            raise Exception("RESEND_API_KEY not set")
        if not EMAIL_SENDER:
            raise Exception("EMAIL_SENDER not set")

        client = ResendClient(api_key=RESEND_API_KEY)
        attachments = []

        if os.path.exists("published_research_paper.docx"):
            with open("published_research_paper.docx", "rb") as f:
                file_data = f.read()
            attachments.append({
                "filename": "published_research_paper.docx",
                "content": base64.b64encode(file_data).decode(),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            })

        client.emails.send({
            "from": EMAIL_SENDER,
            "to": [state["receiver_email"]],
            "subject": f"🔬 Research Report: {state['topic']}",
            "html": state["email_body"],
            "attachments": attachments
        })
        state["log"]("✅ Email sent successfully!")

    except Exception as e:
        state["log"](f"❌ Email failed: {e}")
    return {}

# Build graph
workflow = StateGraph(ResearchState)
for name, fn in [("search", search_published_papers), ("summarize", summarize_papers),
                 ("outline", generate_outline), ("write", write_draft),
                 ("export", export_doc), ("prepare_email", prepare_email), ("send_email", send_email)]:
    workflow.add_node(name, fn)

workflow.set_entry_point("search")
workflow.add_edge("search", "summarize")
workflow.add_edge("summarize", "outline")
workflow.add_edge("outline", "write")
workflow.add_edge("write", "export")
workflow.add_edge("export", "prepare_email")
workflow.add_edge("prepare_email", "send_email")
workflow.add_edge("send_email", END)
graph = workflow.compile()

def run_research_agent(topic: str, receiver_email: str, log_callback=print):
    graph.invoke({
        "topic": topic,
        "receiver_email": receiver_email,
        "papers": [], "summaries": [], "outline": "",
        "draft": "", "references": [], "email_body": "",
        "log": log_callback
    })
